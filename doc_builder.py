import os
import re
import asyncio
import logging
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

import bank_requisites as br
import company

try:  # настройки компании лежат в SQLite бота; без неё работают значения по умолчанию
    from memory import get_setting as _setting
except Exception:  # pragma: no cover
    _setting = None

logger = logging.getLogger(__name__)
logger.info(f"openpyxl version: {openpyxl.__version__}")


# ─── Контроль заполнения документов ────────────────────────────────────────
# Если в готовом документе остался незамещённый {{ПЛЕЙСХОЛДЕР}} или
# плейсхолдер из шаблона получил пустое значение — документ не выдаётся.
# Причина: договор с незаполненными паспортными данными стороны — это дыра
# в комплаенсе, хуже любой неточной формулировки.
# Отключается переменной окружения STRICT_PLACEHOLDERS=0 (тогда только warning).
STRICT_PLACEHOLDERS = os.environ.get("STRICT_PLACEHOLDERS", "1") != "0"

# Плейсхолдеры, которым разрешено быть пустыми (необязательные реквизиты).
ALLOW_EMPTY_PLACEHOLDERS = {
    "{{ЦВЕТ}}",
    "{{НОМ_ТПО}}", "{{ДЕНЬ_ТПО}}", "{{МЕС_ТПО}}", "{{ГОД_ТПО}}",
    "{{ПАСПОРТ_КОД}}",
    "{{БАНК_ПОЛ_СТРОКА2}}", "{{БАНК_КОРР_СТРОКА3}}",
    "{{КОММЕНТАРИЙ}}",
}

# Человекочитаемые названия — чтобы в сообщении Александре было понятно,
# что именно доввести в журнал.
PLACEHOLDER_LABELS = {
    "{{ПОКУПАТЕЛЬ_ФИО}}":           "ФИО покупателя",
    "{{ПОКУПАТЕЛЬ_ДАТА_РОЖДЕНИЯ}}": "дата рождения покупателя",
    "{{ПОКУПАТЕЛЬ_АДРЕС}}":         "адрес покупателя",
    "{{ПОКУПАТЕЛЬ_ИНИЦИАЛЫ}}":      "инициалы покупателя",
    "{{ПАСПОРТ_СЕРИЯ}}":            "серия паспорта",
    "{{ПАСПОРТ_НОМЕР}}":            "номер паспорта",
    "{{ПАСПОРТ_ВЫДАН}}":            "кем выдан паспорт",
    "{{ПАСПОРТ_ДАТА_ВЫДАЧИ}}":      "дата выдачи паспорта",
    "{{ПРОДАВЕЦ_ФИО}}":             "ФИО продавца",
    "{{ПРОДАВЕЦ_ДАТА_РОЖДЕНИЯ}}":   "дата рождения продавца",
    "{{ПРОДАВЕЦ_АДРЕС}}":           "адрес продавца",
    "{{ПРОДАВЕЦ_ИНИЦИАЛЫ}}":        "инициалы продавца",
    "{{ПРОДАВЕЦ_ID_НОМЕР}}":        "номер ID-карты продавца",
    "{{ПРОДАВЕЦ_ID_ВЫДАНА}}":       "кем выдана ID-карта продавца",
    "{{ПРОДАВЕЦ_ID_ДАТА}}":         "дата выдачи ID-карты продавца",
    "{{МАРКА_МОДЕЛЬ}}":             "марка и модель авто",
    "{{VIN}}":                      "VIN",
    "{{ГОД_ВЫП}}":                  "год выпуска",
    "{{ЦЕНА_ЦИФРАМИ}}":             "цена цифрами",
    "{{ЦЕНА_ПРОПИСЬЮ}}":            "цена прописью",
    "{{СУММА_НАЛИЧНЫМИ}}":          "сумма наличными",
    "{{СУММА_НАЛИЧНЫМИ_ПРОПИСЬЮ}}": "сумма наличными прописью",
    "{{ВАЛЮТА_НАЛИЧНЫМИ}}":         "валюта наличных",
    "{{КУРС_ДОЛЛАРА}}":             "курс доллара",
    "{{КУРС_ПОРУЧЕНИЯ}}":           "курс на дату поручения (колонка «Курс USD/RUB»)",
    "{{КУРС_ФАКТИЧЕСКИЙ}}":         "фактический курс конвертации (колонка «Фактический курс»)",
    "{{СУММА_ПОРУЧЕНИЯ}}":          "расчётная сумма поручения в валюте",
    "{{СУММА_ПОРУЧЕНИЯ_ПРОПИСЬЮ}}": "расчётная сумма поручения прописью",
    "{{ДАТА_ПОСТУПЛЕНИЯ}}":         "дата поступления средств",
    "{{ДАТА_РАСЧЕТА}}":             "дата расчёта с получателем",
    "{{СЧЕТ_НОМЕР}}":               "номер счёта",
}


# Допуск сходимости «сумма в валюте × курс = цена в рублях», в рублях.
# Суммы округляются до цента, поэтому точного равенства не бывает: реальное
# расхождение — копейки (49 198,55 × 82,80 = 4 073 639,94 при цене 4 073 640).
# Рубля хватает с запасом и при этом ловится любая опечатка в разряде.
AMOUNT_TOLERANCE_RUB = float(os.environ.get("AMOUNT_TOLERANCE_RUB", "1"))


class MissingTemplateError(Exception):
    """Шаблон под тип счёта не найден.

    Раньше в этом случае счёт молча собирался по шаблону ДРУГОГО типа — с ИНН
    другой юрисдикции в шапке и без QR. Такой документ уходил заказчику и в
    банк как настоящий. Лучше не выдать счёт, чем выдать неправильный.
    """


class AmountMismatchError(Exception):
    """
    Суммы сделки не сходятся: наличные × курс ≠ цена в рублях.

    Поднимается из `_fill_template`. Комплект не формируется — цифры в
    поручении, расписке, акте и отчёте обязаны совпадать до копейки, иначе
    документы противоречат друг другу.
    """

    def __init__(self, doc_name: str, cash, rate, expected, actual):
        self.doc_name = doc_name
        self.cash     = cash
        self.rate     = rate
        self.expected = expected   # цена из ДКП
        self.actual   = actual     # наличные × курс
        super().__init__(
            f"{doc_name}: {cash} × {rate} = {actual:,.2f}, "
            f"а цена в договоре {expected:,.2f}"
        )


def _to_float(v):
    """«86 956,52» → 86956.52. Возвращает None, если это не число."""
    s = str(v or "").replace("\u00a0", "").replace(" ", "").replace(",", ".")
    try:
        return float(s)
    except (TypeError, ValueError):
        return None


def order_amount(price_val, rate) -> float | None:
    """
    Сумма в валюте = цена в рублях / курс, округлённая ДО ЦЕНТА.

    Округление до цента, а не до целого доллара: тогда рублёвый эквивалент
    сходится с ценой ДКП с точностью до копеек (49 198,55 × 82,80 =
    4 073 639,94 при цене 4 073 640). Округление до доллара давало бы на
    каждой сделке хвост в 20–80 рублей, который бухгалтерии пришлось бы
    проводить отдельно.

    Физически продавцу отдают целые доллары (49 199) — разница меньше доллара
    в его пользу нигде не отражается и в учёте не участвует.
    """
    price = _to_float(price_val)
    rate  = _to_float(rate)
    if not price or not rate:
        return None
    return round(price / rate, 2)


def _check_amounts(price_val: float, pairs: list, doc_name: str) -> None:
    """
    Проверяет каждую пару «сумма в валюте — курс» ОТДЕЛЬНО:

        |СУММА_ПОРУЧЕНИЯ  × КУРС_ПОРУЧЕНИЯ   − ЦЕНА| ≤ 1 руб.
        |СУММА_НАЛИЧНЫМИ  × КУРС_ФАКТИЧЕСКИЙ − ЦЕНА| ≤ 1 руб.

    Рублёвая цена одна и та же, поэтому каждая пара обязана сходиться к ней
    независимо. Сравнивать сумму из одной пары с курсом из другой НЕЛЬЗЯ —
    именно так бот остановил акт по сделке 050826001: взял 49 800 (по курсу
    поручения 81,80) и умножил на фактический курс 82,80.

    Допуск в 1 рубль с запасом покрывает округление до цента — реальное
    расхождение составляет копейки.

    pairs — список (сумма, курс, подпись). Пустая сумма или курс не ошибка
    арифметики, а отсутствие данных; ими занимается контроль заполнения.
    """
    if not STRICT_PLACEHOLDERS or not price_val:
        return

    for cash_raw, rate_raw, label in pairs:
        cash = _to_float(cash_raw)
        rate = _to_float(rate_raw)
        if not cash or not rate:
            continue
        actual = cash * rate
        if abs(actual - price_val) > AMOUNT_TOLERANCE_RUB:
            raise AmountMismatchError(f"{doc_name} ({label})", cash_raw, rate_raw,
                                      price_val, actual)


def dkp_number_from(data: dict, fallback: str = "") -> str:
    """
    Номер договора купли-продажи = последние 6 знаков VIN.

    Пример: VIN LVGEU76A1TG062177 → 062177.

    Проверка уникальности не нужна: первичный ключ сделки в системе — номер
    агентского договора, номер ДКП на связки не влияет.

    Колонка журнала «Номер ДКП» (ручной ввод) имеет приоритет — на случай,
    если по бумажному оригиналу номер другой. Если ни колонки, ни VIN нет,
    возвращается fallback (номер сделки) — чтобы документ не остался с
    пустым номером.
    """
    manual = str(data.get("Номер ДКП") or data.get("dkp_number") or "").strip()
    if manual:
        return manual
    vin = re.sub(r"[^A-Za-z0-9]", "", str(data.get("car_vin") or data.get("VIN") or ""))
    if len(vin) >= 6:
        return vin[-6:].upper()
    return fallback


class MissingDataError(Exception):
    """
    Документ собран, но часть плейсхолдеров осталась пустой или незамещённой.

    Поднимается из `_fill_template` при STRICT_PLACEHOLDERS=1. Вызывающий код
    (agent.py) ловит её и возвращает пользователю понятный список полей —
    файл при этом НЕ выдаётся и НЕ заливается на Drive.
    """

    def __init__(self, doc_name: str, missing: list[str], leftover: list[str]):
        self.doc_name = doc_name
        self.missing  = missing
        self.leftover = leftover
        parts = []
        if missing:
            parts.append("не заполнены поля: " + ", ".join(missing))
        if leftover:
            parts.append("не распознаны плейсхолдеры: " + ", ".join(leftover))
        super().__init__(f"{doc_name} — " + "; ".join(parts))


def _fmt_num(v) -> str:
    """Форматирует число для документа: 3997500 → «3 997 500», 1234.5 → «1 234,50»."""
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    if v == int(v):
        return f"{int(v):,}".replace(",", " ")
    return f"{v:,.2f}".replace(",", " ").replace(".", ",")


# ─── Сумма прописью (рубли) ────────────────────────────────────────────────

_UNITS = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
_UNITS_F = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
_TEENS = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
          "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
_TENS = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
         "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
_HUNDREDS = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
             "шестьсот", "семьсот", "восемьсот", "девятьсот"]

# (ед.ч., мн.ч. 2-4, мн.ч. 5+, женский род)
_SCALE = [
    ("", "", "", False),
    ("тысяча", "тысячи", "тысяч", True),
    ("миллион", "миллиона", "миллионов", False),
    ("миллиард", "миллиарда", "миллиардов", False),
]


def _plural(n: int, one: str, few: str, many: str) -> str:
    n100 = n % 100
    n10 = n % 10
    if 11 <= n100 <= 14:
        return many
    if n10 == 1:
        return one
    if 2 <= n10 <= 4:
        return few
    return many


def _three_digits_to_words(n: int, feminine: bool = False) -> str:
    words = []
    h, rem = divmod(n, 100)
    if h:
        words.append(_HUNDREDS[h])
    t, u = divmod(rem, 10)
    if t == 1:
        words.append(_TEENS[u])
    else:
        if t:
            words.append(_TENS[t])
        if u:
            words.append((_UNITS_F if feminine else _UNITS)[u])
    return " ".join(words)


def _int_to_words(n: int) -> str:
    """Целое число прописью, без названия валюты. 3997500 → «три миллиона …»."""
    if n == 0:
        return "ноль"
    groups = []
    scale_idx = 0
    while n > 0:
        n, group = divmod(n, 1000)
        if group:
            groups.append((group, scale_idx))
        scale_idx += 1

    parts = []
    for group, idx in reversed(groups):
        one, few, many, feminine = _SCALE[idx]
        parts.append(_three_digits_to_words(group, feminine=feminine))
        if idx > 0:
            parts.append(_plural(group, one, few, many))
    return " ".join(p for p in parts if p)


def amount_to_words_rub(amount) -> str:
    """
    Преобразует сумму в рублях в строку прописью с копейками.
    Пример: 3997500 -> "Три миллиона девятьсот девяносто семь тысяч пятьсот рублей 00 копеек"
    """
    try:
        amount = float(amount)
    except (TypeError, ValueError):
        return ""

    rub = int(amount)
    kop = round((amount - rub) * 100)

    rub_words = _int_to_words(rub)
    rub_words = rub_words[0].upper() + rub_words[1:]
    rub_label = _plural(rub, "рубль", "рубля", "рублей")
    kop_label = _plural(kop, "копейка", "копейки", "копеек")

    return f"{rub_words} {rub_label} {kop:02d} {kop_label}"


def amount_to_words_plain(amount) -> str:
    """
    Сумма прописью БЕЗ названия валюты — для сумм в долларах.

    В шаблонах валюта стоит отдельным плейсхолдером сразу после скобок:
    «{{СУММА_ПОРУЧЕНИЯ}} ({{СУММА_ПОРУЧЕНИЯ_ПРОПИСЬЮ}}) {{ВАЛЮТА_НАЛИЧНЫМИ}}»,
    поэтому слово «рублей»/«долларов» внутрь скобок писать нельзя.

    Дробная часть даётся центами — суммы после округления до цента почти
    всегда дробные, и в расписке пропись обязана совпадать с цифрами:
    49198.55 → «Сорок девять тысяч сто девяносто восемь 55 центов».

    Слово «долларов» внутрь скобок не пишем — оно уже стоит после них
    как {{ВАЛЮТА_НАЛИЧНЫМИ}}, иначе получился бы дубль.
    Ровные суммы печатаются без хвоста: 12500 → «Двенадцать тысяч пятьсот».
    """
    try:
        amount = float(amount)
    except (TypeError, ValueError):
        return ""

    whole = int(amount)
    frac  = round((amount - whole) * 100)
    if frac == 100:            # 0.999 → 1.00, а не «100 центов»
        whole += 1
        frac = 0

    words = _int_to_words(whole)
    words = words[0].upper() + words[1:]
    if frac == 0:
        return words
    return f"{words} {frac:02d} {_plural(frac, 'цент', 'цента', 'центов')}"


# ─── Варианты шаблона счёта ────────────────────────────────────────────────
# Вёрстку счёта правят в Excel, и держать несколько бланков рядом безопаснее,
# чем перезаписывать боевой файл. Вариант — это суффикс в имени шаблона:
#   invoice_template.xlsx     / invoice_template_direct.xlsx     → v1
#   invoice_template_v2.xlsx  / invoice_template_direct_v2.xlsx  → v2
# Какой печатать, решает настройка INVOICE_TEMPLATE (меню «⚙️ Настройки»).

INVOICE_VARIANT_KEY     = "INVOICE_TEMPLATE"
INVOICE_VARIANT_DEFAULT = "v1"


def templates_dir() -> Path:
    return Path(os.environ.get("TEMPLATES_DIR", "./templates"))


def invoice_template_name(variant: str, is_direct: bool) -> str:
    base = "invoice_template_direct" if is_direct else "invoice_template"
    suffix = "" if variant in ("", "v1", None) else f"_{variant}"
    return f"{base}{suffix}.xlsx"


def invoice_variants() -> list:
    """Варианты бланка, реально лежащие в templates/. Вариант считается
    пригодным, только если есть оба шаблона — прямой и корреспондентский:
    выбрать в настройках половину комплекта нельзя."""
    tpl = templates_dir()
    found = set()
    for f in tpl.glob("invoice_template*.xlsx"):
        stem = f.stem
        for base in ("invoice_template_direct", "invoice_template"):
            if stem == base:
                found.add(INVOICE_VARIANT_DEFAULT)
                break
            if stem.startswith(base + "_"):
                found.add(stem[len(base) + 1:])
                break
    complete = [
        v for v in sorted(found)
        if all((tpl / invoice_template_name(v, d)).exists() for d in (True, False))
    ]
    return complete or [INVOICE_VARIANT_DEFAULT]


def invoice_variant() -> str:
    """Текущий вариант: настройка в БД бота → env → v1.
    Если выбранного комплекта в папке нет, берём первый доступный и пишем
    предупреждение — счёт нужен здесь и сейчас, падать из-за настройки нельзя."""
    val = ""
    if _setting:
        try:
            val = (_setting(INVOICE_VARIANT_KEY) or "").strip()
        except Exception as e:  # pragma: no cover
            logger.warning(f"Не прочитать настройку {INVOICE_VARIANT_KEY}: {e}")
    val = val or os.environ.get(INVOICE_VARIANT_KEY, "").strip() or INVOICE_VARIANT_DEFAULT
    available = invoice_variants()
    if val not in available:
        logger.warning(
            f"Вариант шаблона счёта {val!r} не найден в {templates_dir()}; "
            f"беру {available[0]!r} (доступны: {available})"
        )
        val = available[0]
    return val


class DocumentBuilder:

    def __init__(self):
        self.templates_dir = Path(os.environ.get("TEMPLATES_DIR", "./templates"))
        self.output_dir = Path(tempfile.gettempdir()) / "tg_agent_docs"
        self.output_dir.mkdir(exist_ok=True)

    # ─── АГЕНТСКИЙ ДОГОВОР ────────────────────────────────────────────────

    async def build_contract(self, data: dict, number: str, date: str, commission_pct: float = 1.0) -> str:
        # Тип счёта берётся из явного поля (bank_requisites), а не выводится
        # из пустоты банка-корреспондента, как было до 04.09.2026.
        acc_type = br.resolve_account_type(data)
        is_direct = acc_type == br.DIRECT_RF

        if is_direct:
            template = self.templates_dir / "contract_template_direct.docx"
            if not template.exists():
                raise MissingTemplateError(
                    "Не найден contract_template_direct.docx — шаблон агентского "
                    "договора для прямого счёта в РФ. Договор не собран."
                )
        else:
            template = self.templates_dir / "contract_template.docx"

        logger.info(f"Шаблон АГ договора: {template.name} (тип счёта: {acc_type})")

        if template.exists():
            return await self._fill_template(template, data, number, date,
                                             f"АГ_Договор_{number}", commission_pct)
        return await self._generate_contract(data, number, date, commission_pct)

    async def _generate_contract(self, data: dict, number: str, date: str, commission_pct: float = 1.0) -> str:
        doc = Document()
        self._setup_page(doc)

        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(f"АГЕНТСКИЙ ДОГОВОР № {number}")
        r.bold = True; r.font.size = Pt(13)

        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.add_run("на осуществление платежа в пользу третьего лица")

        doc.add_paragraph(f"г. Бишкек «{date[:2]}» {self._month_name(date[3:5])} {date[6:]} г.")
        doc.add_paragraph()

        buyer = data.get("buyer_name", data.get("company_name", "____________"))
        doc.add_paragraph(
            f"ОсОО «Авто Континент», именуемое в дальнейшем «Агент», в лице "
            f"Генерального директора Колотовкина Ильи Валерьевича, действующего на основании Устава, "
            f"с одной стороны, и {buyer}, именуемый(ая) в дальнейшем «Принципал», "
            f"с другой стороны, заключили настоящий Агентский договор о нижеследующем:"
        )

        sections = [
            ("1. ПРЕДМЕТ ДОГОВОРА", [
                "1.1. Агент обязуется за вознаграждение совершить от своего имени, но за счёт "
                "Принципала действия по передаче денежных средств продавцу транспортного средства.",
                "1.2. Принципал перечисляет денежные средства Агенту безналичным путём, "
                "после чего Агент передаёт их Получателю наличными.",
            ]),
            ("2. ВОЗНАГРАЖДЕНИЕ АГЕНТА", [
                f"2.1. Вознаграждение Агента составляет {commission_pct}% от суммы перевода.",
                "2.2. Вознаграждение уплачивается одновременно с перечислением основной суммы.",
            ]),
            ("3. ОТВЕТСТВЕННОСТЬ СТОРОН", [
                "3.1. Стороны несут ответственность в соответствии с законодательством КР.",
                "3.2. Агент не несёт ответственности за качество приобретаемого ТС.",
            ]),
            ("4. РЕКВИЗИТЫ СТОРОН", [
                "Агент: ОсОО «Авто Континент», ИНН: 01905202610324, "
                "г. Бишкек, Октябрьский район, ул. Матросова, д. 58, Неж.Пом. 2",
                f"Принципал: {buyer}",
            ]),
        ]

        for title, items in sections:
            h = doc.add_paragraph()
            h.add_run(title).bold = True
            for item in items:
                doc.add_paragraph(item)
        doc.add_paragraph()
        self._add_signature_table(doc)

        path = self.output_dir / f"АГ_Договор_{number}.docx"
        doc.save(str(path))
        return str(path)

    # ─── ДКП ТС ───────────────────────────────────────────────────────────

    async def build_dkp(self, data: dict, number: str, date: str) -> str:
        """
        number — номер СДЕЛКИ (агентского договора). Сам номер ДКП внутри
        документа считается по VIN (см. dkp_number_from), а имя файла берётся
        по номеру сделки — чтобы файлы комплекта лежали рядом в папке Drive.

        date — дата сделки; дата самого ДКП подставляется в _fill_template из
        колонки журнала «Дата ДКП», если она заполнена.
        """
        template = self.templates_dir / "dkp_template.docx"
        if template.exists():
            return await self._fill_template(template, data, number, date, f"ДКП_ТС_{number}")
        return await self._generate_dkp(data, number, date)

    async def _generate_dkp(self, data: dict, number: str, date: str) -> str:
        doc = Document()
        self._setup_page(doc)

        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Номер ДКП — по VIN, а не номер сделки (см. dkp_number_from)
        r = t.add_run(
            f"ДОГОВОР КУПЛИ-ПРОДАЖИ ТРАНСПОРТНОГО СРЕДСТВА № "
            f"{dkp_number_from(data, number)}"
        )
        r.bold = True; r.font.size = Pt(13)

        doc.add_paragraph(f"«{date[:2]}» {self._month_name(date[3:5])} {date[6:]} г. г. Бишкек")
        doc.add_paragraph()

        seller   = data.get("seller_name", "____________")
        buyer    = data.get("buyer_name", data.get("company_name", "____________"))
        car      = data.get("car_model", "____________")
        vin      = data.get("car_vin", "____________")
        year     = data.get("car_year", "____")
        color    = data.get("car_color", "____________")
        price    = data.get("car_price", "____________")
        currency = data.get("currency", "RUB")

        doc.add_paragraph(
            f"Гражданин(ка) Кыргызской Республики {seller}, именуемый(ая) в дальнейшем «Продавец», "
            f"с одной стороны, и {buyer}, именуемый(ая) в дальнейшем «Покупатель», "
            f"с другой стороны, заключили настоящий Договор о нижеследующем:"
        )

        items = [
            f"1. Продавец передаёт в собственность Покупателя транспортное средство:\n"
            f"   Марка, модель: {car};\n"
            f"   Идентификационный номер (VIN): {vin};\n"
            f"   Год выпуска: {year};\n"
            f"   № кузова: {vin};\n"
            f"   Цвет: {color}.",
            f"2. Стоимость ТС составляет: {price} {currency}.",
            f"3. Со слов Продавца ТС никому не продано, не заложено, под арестом не состоит.",
            f"4. Покупатель производит оплату через платёжного агента — "
            f"ОсОО «Авто Континент» (ИНН: 01905202610324) — "
            f"в соответствии с Агентским договором № {number} от «{date}».",
            f"5. Право собственности переходит к Покупателю с момента подписания Договора.",
            f"6. Договор составлен в трёх экземплярах.",
        ]

        for item in items:
            doc.add_paragraph(item)
        doc.add_paragraph()
        self._add_signature_table(doc)

        path = self.output_dir / f"ДКП_ТС_{number}.docx"
        doc.save(str(path))
        return str(path)

    # ─── СЧЁТ (XLSX) ──────────────────────────────────────────────────────

    async def build_invoice(self, data: dict, number: str, date: str, commission_pct: float = 1.0) -> str:
        """
        Формирует счёт на оплату.

        Шаблон выбирается по двум признакам: тип счёта (прямой РФ или через
        банк-корреспондент — у них разная шапка и разный ИНН) и вариант
        вёрстки из настройки INVOICE_TEMPLATE.

        Заполнение идёт ТОЛЬКО по плейсхолдерам {{...}}, без привязки к
        координатам ячеек: вёрстку счёта правят в Excel, и любой сдвиг строк
        раньше ломал сборку. Плейсхолдер сам приходит туда, куда его
        поставили в шаблоне.
        """
        bank = br.normalize(data)
        acc_type = bank["account_type"]
        is_direct = acc_type == br.DIRECT_RF
        comp = company.card(_setting)
        variant = invoice_variant()

        logger.info(f"build_invoice: number={number}, тип счёта={acc_type}, "
                    f"вариант шаблона={variant}, "
                    f"банк={bank['bank_name']!r}, БИК={bank['bank_bic']!r}, "
                    f"счёт={bank['account_number']!r}")

        # Никакой подмены на шаблон другого типа: у них разный ИНН в шапке.
        template = self.templates_dir / invoice_template_name(variant, is_direct)
        if not template.exists():
            raise MissingTemplateError(
                f"Не найден {template.name} — шаблон счёта для типа "
                f"«{br.ACCOUNT_TYPE_LABELS[acc_type]}», вариант «{variant}». "
                f"Счёт не собран."
            )

        logger.info(f"Шаблон счёта: {template.name}, "
                    f"размер: {template.stat().st_size} байт")
        wb = openpyxl.load_workbook(str(template))
        ws = wb.active
        logger.info(f"Шаблон загружен, изображений: {len(ws._images)}")

        price_str = str(data.get("car_price", "0")).replace(" ", "").replace(",", ".")
        try:
            price_val = float(price_str)
        except Exception:
            price_val = 0.0

        commission = round(price_val * commission_pct / 100, 2)
        total       = round(price_val + commission, 2)
        currency    = data.get("currency", "RUB")
        acc_cur     = bank["account_currency"] or currency
        buyer       = data.get("buyer_name", data.get("company_name", ""))
        car         = (f"{data.get('car_model', '')} год выпуска {data.get('car_year', '')} "
                       f"VIN {data.get('car_vin', '')}").strip()

        day_n  = date[0:2]
        mon_n  = date[3:5]
        year_n = date[6:10]
        date_str = f"{day_n} {self._month_name(mon_n)} {year_n}"

        total_fmt   = f"{total:,.2f}".replace(",", " ")
        total_words = amount_to_words_rub(total) if acc_cur == "RUB" else ""

        # Шапка счёта: данные компании из карточки, данные счёта из профиля.
        # Какой из двух ИНН печатать, решает не код, а шаблон — их два, и
        # каждый знает свою юрисдикцию. Старые имена плейсхолдеров оставлены
        # синонимами, чтобы не обновлённый шаблон не превратился в пустой бланк.
        replacements = dict(company.placeholders(_setting))
        replacements.update(br.placeholders(data))
        replacements.update({
            "{{QR}}":               "",   # заменяется изображением ниже
            "{{QR_CODE}}":          "",
            "{{ACCOUNT_NUMBER}}":   bank["account_number"],
            "{{ПОЛУЧАТЕЛЬ}}":       comp["company_name"],
            "{{BANK_BEN_NAME2}}":   comp["company_name"],
            "{{ПОЛУЧАТЕЛЬ_ИНН}}":   comp["company_inn_rf"] if is_direct else comp["company_inn"],
            "{{BANK_BEN_INN}}":     comp["company_inn_rf"] if is_direct else comp["company_inn"],
            "{{ПОЛУЧАТЕЛЬ_КПП}}":   comp["company_kpp_rf"] if is_direct else "",
            # Тело счёта — раньше эти строки писались по координатам B14/G20/D23…
            "{{НОМЕР}}":            number,
            "{{ДАТА}}":             f"{date_str} г.",
            "{{ДАТА_ЦИФРАМИ}}":     date,
            "{{ПОКУПАТЕЛЬ}}":       buyer,
            "{{АВТО}}":             car,
            "{{ВАЛЮТА}}":           acc_cur,
            "{{СЧЕТ_ВАЛЮТА}}":      acc_cur,
            "{{ИТОГО_ЧИСЛОМ}}":     total_fmt,
            "{{ИТОГО_ПРОПИСЬЮ}}":   total_words,
            "{{КОЛИЧЕСТВО}}":       "1",
        })
        if is_direct:
            replacements.update({
                "{{BANK_DIRECT_NAME}}": bank["bank_name"],
                "{{BANK_DIRECT_BIK}}":  bank["bank_bic"],
                "{{BANK_DIRECT_CORR}}": bank["bank_corr_acc"],
                "{{BANK_DIRECT_KPP}}":  comp["company_kpp_rf"],
            })
        else:
            replacements.update({
                "{{BANK_CORR_NAME}}": bank["corr_bank_name"],
                "{{BANK_CORR_BIK}}":  bank["corr_bank_bic"],
                "{{BANK_CORR_ACC}}":  bank["corr_bank_acc"],
                "{{BANK_BEN_NAME}}":  bank["bank_name"],
                "{{BANK_BEN_BIK}}":   bank["bank_bic"],
                "{{BANK_BEN_CORR}}":  bank["bank_corr_acc"],
                "{{BANK_BEN_LINE2}}": (f"БИК: {bank['bank_bic']}, корр. счёт: "
                                       f"{bank['bank_corr_acc']}")
                                      if bank["bank_bic"] else "",
            })

        # Суммы записываем числами, иначе формулы «Сумма» и «Итого» в
        # шаблоне посчитают по тексту и дадут ноль.
        numeric = {
            "{{ЦЕНА_АВТО}}":       price_val,
            "{{КОМИССИЯ_СУММА}}":  commission,
            "{{ИТОГО_ЧИСЛОМ_ЧИСЛО}}": total,
        }

        # Координату ячейки с QR запоминаем до замены — после она станет пустой.
        qr_cell_coord = None
        if is_direct:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value in ("{{QR}}", "{{QR_CODE}}"):
                        qr_cell_coord = cell.coordinate
                        break

        # ── Заполнение по плейсхолдерам ───────────────────────────────────
        # iter_rows отдаёт объединённые ячейки как MergedCell со значением
        # None, поэтому фильтр по строке заодно защищает от записи в них.
        filled = 0
        for row in ws.iter_rows():
            for c in row:
                if not isinstance(c.value, str):
                    continue
                raw = c.value.strip()
                if raw in numeric:
                    c.value = numeric[raw]
                    filled += 1
                    continue
                val = c.value
                for ph, rep in replacements.items():
                    if ph in val:
                        val = val.replace(ph, str(rep))
                if val != c.value:
                    c.value = val
                    filled += 1

        leftover = sorted({
            m for row in ws.iter_rows() for c in row
            if isinstance(c.value, str)
            for m in re.findall(r"\{\{[^}]+\}\}", c.value)
        })
        logger.info(f"Счёт: заполнено ячеек {filled}, вариант шаблона {variant}")
        if leftover:
            # Не роняем выдачу: счёт с лишним плейсхолдером видно глазом,
            # а вот молчаливо не собранный счёт хуже.
            logger.warning(f"В счёте остались незаполненные плейсхолдеры: {leftover}")

        # ── QR код для прямого счёта ──────────────────────────────────────
        if is_direct and qr_cell_coord:
            try:
                import qrcode
                from qrcode.constants import ERROR_CORRECT_M
                from openpyxl.drawing.image import Image as XLImage
                import io

                account   = bank["account_number"]
                bic       = bank["bank_bic"]
                corr      = bank["bank_corr_acc"]
                bank_name = bank["bank_name"]
                # QR по ГОСТ Р 56042 читает банковское приложение плательщика,
                # поэтому здесь всегда российские реквизиты компании.
                kpp       = comp["company_kpp_rf"]
                payee     = comp["company_name"]
                payee_inn = comp["company_inn_rf"]
                sum_kopecks = int(round(total * 100))

                # ГОСТ Р 56042 — порядок обязательных полей:
                # ST00012 → Name → PersonalAcc → BankName → BIC → CorrespAcc
                # Затем опциональные: PayeeINN, KPP, Sum, Purpose
                car_vin = data.get("car_vin", "")
                qr_str = (
                    "ST00012|"
                    f"Name={payee}|"
                    f"PersonalAcc={account}|"
                    f"BankName={bank_name}|"
                    f"BIC={bic}|"
                    f"CorrespAcc={corr}|"
                    f"PayeeINN={payee_inn}|"
                    f"KPP={kpp}|"
                    f"Sum={sum_kopecks}|"
                    f"Purpose=Оплата по счету №{number} от {date} по Агентскому договору №{number} от {date} за автомобиль VIN {car_vin}. Без НДС."
                )

                # Кодируем в cp1251 как требует ГОСТ
                qr_bytes = qr_str.encode("cp1251")

                qr = qrcode.QRCode(
                    version=None,
                    error_correction=ERROR_CORRECT_M,
                    box_size=4,
                    border=2,
                )
                qr.add_data(qr_bytes, optimize=0)
                qr.make(fit=True)
                img = qr.make_image(fill_color="black", back_color="white")

                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)

                xl_img = XLImage(buf)
                xl_img.width  = 90
                xl_img.height = 90
                ws.add_image(xl_img, qr_cell_coord)
                logger.info(f"QR код вставлен в {qr_cell_coord} ({len(qr_bytes)} байт)")
                logger.info(f"QR содержимое: {qr_str}")
            except ImportError:
                logger.warning("Библиотека qrcode не установлена — QR пропущен")
            except Exception as e:
                logger.warning(f"Ошибка генерации QR: {e}", exc_info=True)

        # Настройка области печати
        try:
            from openpyxl.worksheet.properties import WorksheetProperties, PageSetupProperties
            if ws.sheet_properties is None:
                ws.sheet_properties = WorksheetProperties()
            if ws.sheet_properties.pageSetUpPr is None:
                ws.sheet_properties.pageSetUpPr = PageSetupProperties()
            ws.print_area = ws.dimensions
            ws.page_setup.orientation = "portrait"
            ws.page_setup.fitToPage = True
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0
        except Exception as e:
            logger.warning(f"Не удалось настроить page_setup для PDF: {e}")

        path = self.output_dir / f"Счёт_{number}.xlsx"
        wb.save(str(path))

        wb_check = openpyxl.load_workbook(str(path))
        n_images = len(wb_check.active._images)
        logger.info(f"Счёт сохранён, изображений в файле: {n_images}")

        if n_images == 0 and template.exists():
            logger.warning("Изображения потерялись — восстанавливаю из шаблона")
            self._restore_images_from_template(path, template)

        return str(path)

    # ─── АКТ ВЫПОЛНЕННЫХ УСЛУГ ────────────────────────────────────────────

    async def build_act(self, data: dict, contract_number: str, contract_date: str,
                        act_date: str, commission_pct: float = 1.0) -> str:
        """
        Формирует акт об оказании услуг по агентскому договору.

        contract_date — дата АГ договора/счёта/ДКП (ДД.ММ.ГГГГ). По решению
        пользователя во все три места («от … г.») подставляется одна и та же
        дата — дата основного договора. Плейсхолдеры в шаблоне разведены
        (_ДОГОВОРА / _СЧЕТА / _ДКП) — если когда-нибудь потребуется различать,
        менять только эту функцию, шаблон уже готов.

        act_date — дата самого акта (дата закрывающего платежа).
        """
        template = self.templates_dir / "act_template.docx"
        if not template.exists():
            raise FileNotFoundError(
                f"Шаблон акта не найден: {template}. "
                "Положите act_template.docx в папку templates/."
            )

        a_day, a_month, a_year = self._date_parts(act_date)

        # Перекрываем только дату акта — она единственная отличается от даты
        # сделки. Номера всех документов (_АКТА/_ДОГОВОРА/_СЧЕТА/_ДКП), даты
        # договора/счёта/ДКП, суммы, комиссия и итог — из базового набора
        # _fill_template, там же они считаются для отчёта и счёта.
        extra = {
            "{{ДЕНЬ_АКТА}}":  a_day,
            "{{МЕСЯЦ_АКТА}}": a_month,
            "{{ГОД_АКТА}}":   a_year,
        }

        return await self._fill_template(
            template, data, contract_number, contract_date,
            f"Акт_{contract_number}", commission_pct,
            extra_replacements=extra,
        )

    # ─── ОТЧЁТ АГЕНТА ─────────────────────────────────────────────────────

    async def build_report(self, data: dict, contract_number: str, contract_date: str,
                           report_date: str, received_date: str, settlement_date: str,
                           commission_pct: float = 1.0) -> str:
        """
        Формирует отчёт агента по агентскому договору.

        contract_date    — дата АГ договора / счёта / ДКП (ДД.ММ.ГГГГ).
        report_date      — дата самого отчёта (по умолчанию = дата акта).
        received_date    — дата зачисления средств от принципала ({{ДАТА_ПОСТУПЛЕНИЯ}}).
        settlement_date  — дата передачи наличных получателю ({{ДАТА_РАСЧЕТА}}).

        Все три даты по умолчанию берутся из даты закрывающего платежа
        (см. build_report_impl в agent.py) — отчёт, акт и расчёт обычно
        оформляются одним днём.
        """
        template = self.templates_dir / "otchet_agenta_template.docx"
        if not template.exists():
            raise FileNotFoundError(
                f"Шаблон отчёта агента не найден: {template}. "
                "Положите otchet_agenta_template.docx в папку templates/."
            )

        r_day, r_month, r_year = self._date_parts(report_date)

        # Даты поступления и расчёта идут в шаблон одним полем ДД.ММ.ГГГГ.
        # Кладём их и в data — базовый набор _fill_template читает их оттуда,
        # чтобы значение было одно и то же, откуда бы отчёт ни собирали.
        data = dict(data)
        data["payment_received_date"] = received_date
        data["settlement_date"]       = settlement_date

        extra = {
            "{{ДЕНЬ_ОТЧЕТА}}":  r_day,
            "{{МЕСЯЦ_ОТЧЕТА}}": r_month,
            "{{ГОД_ОТЧЕТА}}":   r_year,
        }

        return await self._fill_template(
            template, data, contract_number, contract_date,
            f"Отчёт_агента_{contract_number}", commission_pct,
            extra_replacements=extra,
        )

    # ─── РАСПИСКА О ПОЛУЧЕНИИ ДЕНЕЖНЫХ СРЕДСТВ ───────────────────────────

    async def build_receipt(self, data: dict, contract_number: str, contract_date: str,
                            receipt_date: str, commission_pct: float = 1.0) -> str:
        """
        Формирует расписку продавца о получении наличных денежных средств
        (Приложение № 1 к акту об оказании услуг).

        contract_date — дата АГ договора / ДКП (ДД.ММ.ГГГГ). Подставляется во все
        места «от «..» ... г.», где упоминаются номера договоров — так же как в акте.

        receipt_date — дата самой расписки. Совпадает с датой акта, то есть с датой
        хронологически последнего платежа по сделке.
        """
        template = self.templates_dir / "raspiska_template.docx"
        if not template.exists():
            raise FileNotFoundError(
                f"Шаблон расписки не найден: {template}. "
                "Положите raspiska_template.docx в папку templates/."
            )

        n = self._normalize

        # Разбор даты «ДД.ММ.ГГГГ» → (день, месяц-словом, год)
        def _parts(date_str):
            date_str = (date_str or "").strip()
            if len(date_str) < 10:
                return "", "", ""
            return date_str[0:2], self._month_name(date_str[3:5]), date_str[6:10]

        r_day, r_month, r_year = _parts(receipt_date)

        # ── Блок получателя денег — продавец (гражданин КР) ───────────────
        receiver = n(data.get("seller_full_details", "")).strip()
        if not receiver:
            bits = [n(data.get("seller_name", ""))]
            if data.get("seller_birth_date"):
                bits.append(f"{data['seller_birth_date']} года рождения")
            id_num = data.get("seller_id_number", "") or data.get("seller_id", "")
            if id_num:
                id_str = f"ID-карта № {id_num}"
                if data.get("seller_id_issued_by"):
                    id_str += f", выдана {n(data['seller_id_issued_by'])}"
                if data.get("seller_id_issued_date"):
                    id_str += f" {data['seller_id_issued_date']}"
                bits.append(id_str)
            if data.get("seller_address"):
                bits.append(f"зарегистрирован(а) по адресу: {n(data['seller_address'])}")
            receiver = ", ".join(b for b in bits if b)

        receiver_initials = (n(data.get("seller_initials", "")).strip()
                             or n(data.get("seller_name", "")).strip())

        # ── Сумма в рублях прописью ───────────────────────────────────────
        # В журнале поле может быть пустым — считаем сами, чтобы в расписке
        # никогда не оставалось незамещённого плейсхолдера.
        price_words = (data.get("car_price_words", "") or "").strip()
        if not price_words:
            price_str = str(data.get("car_price", "0")).replace(" ", "").replace(",", ".")
            try:
                price_words = amount_to_words_rub(float(price_str))
            except (TypeError, ValueError):
                price_words = ""
            # В шаблоне слово «рублей» стоит уже после скобок, поэтому отбрасываем
            # хвост «рублей NN копеек» — иначе получится «... рублей 00 копеек) рублей РФ».
            cut = re.search(r"\s+рубл(ь|я|ей)\b", price_words)
            if cut:
                price_words = price_words[:cut.start()]

        extra = {
            # Дата расписки = дата акта (дата закрывающего платежа)
            "{{ДЕНЬ_АКТА}}":  r_day,
            "{{МЕСЯЦ_АКТА}}": r_month,
            "{{ГОД_АКТА}}":   r_year,

            # Получатель наличных — продавец
            "{{ПОЛУЧАТЕЛЬ_БЛОК}}":     receiver,
            "{{ПОЛУЧАТЕЛЬ_ИНИЦИАЛЫ}}": receiver_initials,

            # Рублёвый эквивалент прописью
            "{{ЦЕНА_ПРОПИСЬЮ}}": price_words,
        }
        # {{НОМЕР}}, {{ДЕНЬ}}/{{МЕСЯЦ}}/{{ГОД}}, {{СУММА_НАЛИЧНЫМИ*}},
        # {{ВАЛЮТА_НАЛИЧНЫМИ}}, {{КУРС_ДОЛЛАРА}}, {{ЦЕНА_ЦИФРАМИ}},
        # {{МАРКА_МОДЕЛЬ}}, {{VIN}} — покрыты базовым набором в _fill_template.

        return await self._fill_template(
            template, data, contract_number, contract_date,
            f"Расписка_{contract_number}", commission_pct,
            extra_replacements=extra,
        )

    # ─── ВОССТАНОВЛЕНИЕ КАРТИНОК ИЗ ШАБЛОНА xlsx ─────────────────────────

    def _restore_images_from_template(self, output_path: Path, template_path: Path):
        """
        Запасной путь: копирует печать/подпись (xl/media/*, xl/drawings/*) и связи
        из шаблона прямо в zip-архив готового файла, на случай если openpyxl
        потерял изображения при load/save.
        """
        import zipfile
        import shutil
        import re

        tmp_path = output_path.with_suffix(".tmp.xlsx")

        with zipfile.ZipFile(template_path, "r") as tz:
            template_names = set(tz.namelist())
            media_files   = [n for n in template_names if n.startswith("xl/media/")]
            drawing_files = [n for n in template_names if n.startswith("xl/drawings/")]

            if not media_files:
                logger.warning("В шаблоне нет xl/media/* — восстановление невозможно")
                return

            with zipfile.ZipFile(output_path, "r") as oz:
                out_names = set(oz.namelist())

                with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as nz:
                    # копируем всё содержимое готового файла, кроме того что заменим
                    skip = set(media_files) | set(drawing_files)
                    for item in oz.infolist():
                        if item.filename in skip:
                            continue
                        data = oz.read(item.filename)
                        if item.filename == "xl/worksheets/sheet1.xml":
                            # добавляем ссылку <drawing r:id="..."/> перед </worksheet>, если её нет
                            text = data.decode("utf-8")
                            if "<drawing " not in text:
                                if 'xmlns:r=' not in text.split('>', 1)[0]:
                                    text = text.replace(
                                        "<worksheet xmlns=",
                                        '<worksheet xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns=',
                                        1,
                                    )
                                text = text.replace("</worksheet>", '<drawing r:id="rIdDrawing1"/></worksheet>')
                                data = text.encode("utf-8")
                        elif item.filename == "xl/worksheets/_rels/sheet1.xml.rels":
                            text = data.decode("utf-8")
                            if "drawing1.xml" not in text:
                                text = text.replace(
                                    "</Relationships>",
                                    '<Relationship Id="rIdDrawing1" '
                                    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing" '
                                    'Target="../drawings/drawing1.xml"/></Relationships>'
                                )
                                data = text.encode("utf-8")
                        elif item.filename == "[Content_Types].xml":
                            text = data.decode("utf-8")
                            additions = ""
                            if "PartName=\"/xl/drawings/drawing1.xml\"" not in text:
                                additions += ('<Override PartName="/xl/drawings/drawing1.xml" '
                                               'ContentType="application/vnd.openxmlformats-officedocument.drawing+xml"/>')
                            if "Extension=\"png\"" not in text:
                                additions += '<Default Extension="png" ContentType="image/png"/>'
                            if additions:
                                text = text.replace("</Types>", additions + "</Types>")
                                data = text.encode("utf-8")
                        nz.writestr(item, data)

                    if "xl/worksheets/_rels/sheet1.xml.rels" not in out_names:
                        # на случай если у листа вообще не было _rels
                        rels = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                                '<Relationship Id="rIdDrawing1" '
                                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing" '
                                'Target="../drawings/drawing1.xml"/></Relationships>')
                        nz.writestr("xl/worksheets/_rels/sheet1.xml.rels", rels)

                    # копируем media и drawings из шаблона как есть
                    for name in media_files + drawing_files:
                        nz.writestr(name, tz.read(name))

        shutil.move(str(tmp_path), str(output_path))

        # финальная проверка
        try:
            wb_final = openpyxl.load_workbook(str(output_path))
            logger.info(f"После восстановления изображений: {len(wb_final.active._images)}")
        except Exception as e:
            logger.error(f"Ошибка проверки файла после восстановления картинок: {e}", exc_info=True)

    # ─── КОНВЕРТАЦИЯ В PDF ────────────────────────────────────────────────

    # ─── КАРТОЧКА КОМПАНИИ ДЛЯ КОНТРАГЕНТА ───────────────────────────────

    async def build_company_card(self, profile: dict, profile_name: str,
                                 getter=None) -> str:
        """Карточка предприятия: данные компании плюс реквизиты одного счёта.

        Собирается из обычного шаблона templates/company_card_template.docx —
        как и остальные документы, чтобы вид карточки правился в Word, а не в
        коде. Строки таблицы, значение которых оказалось пустым, удаляются:
        так у прямого российского счёта сам собой исчезает блок
        банка-корреспондента, а незаполненные поля карточки не оставляют
        пустых строк с прочерком.
        """
        template = self.templates_dir / "company_card_template.docx"
        if not template.exists():
            raise MissingTemplateError(
                "Не найден company_card_template.docx — шаблон карточки "
                "предприятия. Карточка не собрана."
            )

        values = dict(company.placeholders(getter))
        values.update(br.placeholders(profile))

        doc = Document(str(template))

        def fill(text: str) -> str:
            for ph, val in values.items():
                if ph in text:
                    text = text.replace(ph, str(val))
            return text

        def replace_in(par):
            """Подстановка с запасом на разорванный плейсхолдер.

            Сначала пробуем по одному run — так сохраняется форматирование.
            Если после этого «{{» в абзаце ещё остались, значит Word разрезал
            плейсхолдер между runs при ручной правке шаблона: тогда сливаем
            абзац в первый run.
            """
            for run in par.runs:
                if "{{" in run.text:
                    run.text = fill(run.text)
            if "{{" in par.text and par.runs:
                merged = fill("".join(r.text for r in par.runs))
                par.runs[0].text = merged
                for extra in par.runs[1:]:
                    extra.text = ""

        for par in doc.paragraphs:
            replace_in(par)

        for table in doc.tables:
            drop = []
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        replace_in(par)
                # Пустое значение — строка не нужна.
                if len(row.cells) == 2 and not row.cells[1].text.strip():
                    drop.append(row)
            for row in drop:
                row._element.getparent().remove(row._element)

        safe = re.sub(r"[^\w\-]+", "_", profile_name, flags=re.UNICODE).strip("_")
        path = self.output_dir / f"Карточка_предприятия_{safe}.docx"
        doc.save(str(path))
        logger.info(f"Карточка предприятия собрана: {path.name}")
        return str(path)

    async def convert_to_pdf(self, filepath: str) -> str | None:
        """
        Конвертирует файл в PDF через LibreOffice.
        ИСПРАВЛЕНО: asyncio.create_subprocess_exec вместо subprocess.run —
        не блокирует event loop Telegram-бота.
        Возвращает путь к PDF или None если LibreOffice недоступен / ошибка.
        """
        try:
            proc = await asyncio.create_subprocess_exec(
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(self.output_dir), filepath,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            try:
                stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=60)
                if stderr:
                    logger.debug(f"LibreOffice stderr: {stderr.decode(errors='ignore')}")
            except asyncio.TimeoutError:
                proc.kill()
                await proc.communicate()
                logger.warning("LibreOffice: таймаут конвертации (> 60 сек)")
                return None

            pdf_path = str(filepath).rsplit(".", 1)[0] + ".pdf"
            if Path(pdf_path).exists() and Path(pdf_path).stat().st_size > 0:
                return pdf_path

            logger.warning(f"LibreOffice: PDF не создан для {filepath}")
            return None

        except FileNotFoundError:
            logger.warning("LibreOffice не установлен — PDF конвертация недоступна")
        except Exception as e:
            logger.error(f"Ошибка конвертации PDF: {e}")
        return None

    # ─── ВСПОМОГАТЕЛЬНЫЕ ─────────────────────────────────────────────────

    def _setup_page(self, doc):
        section = doc.sections[0]
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    def _add_signature_table(self, doc):
        table = doc.add_table(rows=3, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 0).text = "Продавец (Агент):"
        table.cell(0, 1).text = "Покупатель (Принципал):"
        table.cell(1, 0).text = "ОсОО «Авто Континент»"
        table.cell(2, 0).text = "\n_________________ / Колотовкин И.В. /"
        table.cell(2, 1).text = "\n_________________ /____________/"

    def _month_name(self, month_num: str) -> str:
        months = {
            "01": "января",  "02": "февраля", "03": "марта",
            "04": "апреля",  "05": "мая",     "06": "июня",
            "07": "июля",    "08": "августа", "09": "сентября",
            "10": "октября", "11": "ноября",  "12": "декабря",
        }
        return months.get(month_num, month_num)

    def _date_parts(self, date_str: str) -> tuple[str, str, str]:
        """«07.08.2026» → («07», «августа», «2026»). Пустая/битая дата → («», «», «»)."""
        date_str = (date_str or "").strip()
        if len(date_str) < 10:
            return "", "", ""
        return date_str[0:2], self._month_name(date_str[3:5]), date_str[6:10]

    @staticmethod
    def _scan_placeholders(doc) -> set:
        """
        Возвращает множество всех {{ПЛЕЙСХОЛДЕРОВ}}, встречающихся в документе:
        в тексте, таблицах, колонтитулах.

        Используется дважды: до подстановки — чтобы знать, какие поля документ
        вообще требует, и после — чтобы поймать незамещённые.
        """
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        found = set()

        def scan_para(para):
            text = "".join(
                t.text or ""
                for r in para._element.findall(f"{{{W}}}r")
                for t in r.findall(f"{{{W}}}t")
            )
            found.update(re.findall(r"\{\{[^{}]+\}\}", text))

        for para in doc.paragraphs:
            scan_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        scan_para(para)
        for section in doc.sections:
            for para in section.header.paragraphs:
                scan_para(para)
            for para in section.footer.paragraphs:
                scan_para(para)
        return found

    @staticmethod
    def _normalize(value: str) -> str:
        """
        Приводит строку к нормальному регистру: первая буква каждого слова — заглавная,
        остальные — строчные. Нужно потому что в документах КР данные часто написаны КАПСОМ.
        Аббревиатуры (МКК, СОМ, УФМС и т.д.) ВСЕГДА приводятся к верхнему регистру,
        даже если строка уже в смешанном регистре (например агент написал "Мкк").
        """
        if not value or not value.strip():
            return value
        if not any(c.isalpha() for c in value):
            return value

        # Аббревиатуры — всегда верхний регистр независимо от входного регистра
        UPPER_WORDS = {
            # Органы КР
            "МКК", "СОМ", "ГУВД", "УВД", "ОВД", "МВД",
            # Органы РФ
            "УФМС", "ОМВД", "ОУФМС", "МФЦ", "ФМС", "ФСБ", "МЧС",
            # Орг. формы
            "ОАО", "ООО", "АО", "ЗАО", "ПАО", "ГУП", "МУП", "НКО", "ИП",
            # Страны и союзы
            "РФ", "КР", "СНГ", "СССР", "США", "ЕС",
            # Прочие
            "ИНН", "КПП", "БИК", "ОГРН", "СНИЛС", "VIN",
            "ТПО", "ДКП", "АГ",
        }

        # Слова со специальным (нестандартным) регистром
        SPECIAL_CASE_WORDS = {
            "ОСОО": "ОсОО",
            "ОсОО": "ОсОО",
        }

        # Предлоги/союзы (не первое слово) оставляем строчными
        lower_words = {"и", "в", "на", "по", "из", "за", "от", "до", "при", "для",
                       "или", "но", "а", "не", "то", "со", "об", "под", "над"}

        # Если строка уже в смешанном регистре — только исправляем аббревиатуры
        already_mixed = any(c.islower() for c in value)

        words = value.split()
        result = []
        for i, word in enumerate(words):
            clean = word.strip('.,;:()«»"\'')
            if clean.upper() in SPECIAL_CASE_WORDS:
                # Слова с нестандартным регистром (например ОсОО)
                prefix = word[:len(word) - len(word.lstrip('.,;:()«»"\''))]
                suffix = word[len(word.rstrip('.,;:()«»"\'')):]
                result.append(prefix + SPECIAL_CASE_WORDS[clean.upper()] + suffix)
            elif clean.upper() in UPPER_WORDS:
                # Аббревиатура — всегда верхний регистр
                # Сохраняем знаки препинания по краям
                prefix = word[:len(word) - len(word.lstrip('.,;:()«»"\''))]
                suffix = word[len(word.rstrip('.,;:()«»"\'')):]
                result.append(prefix + clean.upper() + suffix)
            elif already_mixed:
                # Строка уже нормальная — не трогаем остальные слова
                result.append(word)
            elif any(c.isdigit() for c in word):
                result.append(word[0].upper() + word[1:].lower() if word[0].isalpha() else word)
            elif word.isalpha() and i > 0 and word.lower() in lower_words:
                result.append(word.lower())
            else:
                result.append(word.capitalize())
        return " ".join(result)

    async def _fill_template(self, template_path, data, number, date, output_name,
                              commission_pct: float = 1.0,
                              extra_replacements: dict | None = None) -> str:
        from lxml import etree
        from copy import deepcopy

        doc = Document(str(template_path))

        # Нормализуем регистр текстовых полей — данные из КР-документов часто приходят КАПСОМ
        n = self._normalize


        day   = date[0:2]
        month = date[3:5]
        year  = date[6:10]

        price_str = data.get("car_price", "0").replace(" ", "").replace(",", ".")
        try:
            price_val = float(price_str)
            price_fmt = f"{price_val:,.0f}".replace(",", " ")
        except Exception:
            price_fmt = price_str
            price_val = 0

        # ── Два курса и две суммы в валюте ─────────────────────────────────
        # Рублёвая сумма зафиксирована договором купли-продажи и не меняется.
        # Поручение показывает расчёт на СВОЮ дату (курс поручения), а
        # конвертация происходит позже и по другому курсу — поэтому
        # фактическая сумма в валюте отличается от расчётной. Расхождение
        # прямо описано в текстах акта, отчёта и примечании к поручению,
        # прятать его не нужно, нужно правильно проставить.
        rate_order = _to_float(data.get("exchange_rate"))          # курс поручения
        rate_fact  = _to_float(data.get("Фактический курс")
                               or data.get("exchange_rate_fact"))  # курс конвертации

        # РАСЧЁТНАЯ сумма — только в поручение, в паре с курсом поручения.
        # Колонка журнала «Сумма нал. (USD)» (cash_amount) — это она и есть.
        order_val = _to_float(data.get("cash_amount"))
        if not order_val:
            order_val = order_amount(price_val, rate_order)

        # ФАКТИЧЕСКАЯ сумма — в расписку, акт и отчёт, в паре с фактическим
        # курсом. Считается заново от фактического курса, а не берётся из
        # «Сумма нал.»: та посчитана по курсу поручения, и умножать её на
        # другой курс — значит складывать доллары одного курса с курсом
        # другого. Ровно на этом встал акт по сделке 050826001.
        # Приоритет у колонки «Сумма выдана (USD)»: бот пишет туда значение
        # при первой сборке документа, чтобы повторная генерация не дала
        # другое число, чем в уже подписанной расписке.
        cash_val = _to_float(data.get("Сумма выдана (USD)")
                             or data.get("cash_paid_amount"))
        if not cash_val:
            cash_val = order_amount(price_val, rate_fact)

        cash_fmt  = _fmt_num(cash_val) if cash_val else ""
        order_fmt = _fmt_num(order_val) if order_val else ""

        # Пропись считается по числу, а не берётся из журнала: с округлением
        # до цента ручная пропись почти наверняка разойдётся с цифрами.
        cash_words  = amount_to_words_plain(cash_val) if cash_val else ""
        order_words = amount_to_words_plain(order_val) if order_val else ""

        # ── Комиссия и итог — считаются в одном месте для всех документов ──
        # Комиссия = ЦЕНА_ЦИФРАМИ × КОМИССИЯ% / 100, база — сумма Поручения
        # (car_price). Считаем здесь, а не в каждом build_*, чтобы акт, отчёт
        # и счёт никогда не разошлись в арифметике.
        commission_val = round(price_val * commission_pct / 100, 2)
        total_val      = round(price_val + commission_val, 2)

        # ── Проверка арифметики сделки ─────────────────────────────────────
        # Комиссия и итог считаются здесь же, поэтому сойтись обязаны всегда.
        # А вот каждая пара «сумма в валюте — курс» проверяется отдельно:
        # обе обязаны сходиться к одной и той же рублёвой цене.
        _check_amounts(price_val, [
            (order_val, rate_order, "поручение"),
            (cash_val,  rate_fact,  "расчёт наличными"),
        ], output_name)

        # Реквизиты приводим к модели один раз на документ. Данные компании
        # приходят из карточки (company.py), данные счёта — из сделки.
        bank = br.normalize(data)
        comp = company.placeholders(_setting)

        replacements = {
            "{{НОМЕР}}":   number,
            "{{ДЕНЬ}}":    day,
            "{{МЕСЯЦ}}":   self._month_name(month),
            "{{ГОД}}":     year,
            "{{КОМИССИЯ}}": str(commission_pct),

            # Комиссия и итог
            "{{КОМИССИЯ_ЦИФРАМИ}}":     _fmt_num(commission_val),
            "{{КОМИССИЯ_ПРОПИСЬЮ}}":    amount_to_words_rub(commission_val),
            "{{СУММА_ИТОГО_ЦИФРАМИ}}":  _fmt_num(total_val),
            "{{СУММА_ИТОГО_ПРОПИСЬЮ}}": amount_to_words_rub(total_val),

            # Покупатель (гражданин РФ)
            "{{ПОКУПАТЕЛЬ_ФИО}}":           n(data.get("buyer_name", "")),
            "{{ПОКУПАТЕЛЬ_ДАТА_РОЖДЕНИЯ}}": data.get("buyer_birth_date", ""),
            "{{ПОКУПАТЕЛЬ_АДРЕС}}":         n(data.get("buyer_address", "")),
            "{{ПОКУПАТЕЛЬ_ИНИЦИАЛЫ}}":      n(data.get("buyer_initials", "")),
            "{{ПОКУПАТЕЛЬ_ПОЛНЫЕ_ДАННЫЕ}}": n(data.get("buyer_full_details", data.get("buyer_name", ""))),

            # Паспорт покупателя (РФ)
            "{{ПАСПОРТ_СЕРИЯ}}":       data.get("passport_series", ""),
            "{{ПАСПОРТ_НОМЕР}}":       data.get("passport_number", ""),
            "{{ПАСПОРТ_ВЫДАН}}":       n(data.get("passport_issued_by", "")),
            "{{ПАСПОРТ_КОД}}":         data.get("passport_code", ""),
            "{{ПАСПОРТ_ДАТА_ВЫДАЧИ}}": data.get("passport_issued_date", ""),

            # Продавец (гражданин КР)
            "{{ПРОДАВЕЦ_ФИО}}":           n(data.get("seller_name", "")),
            "{{ПРОДАВЕЦ_ДАТА_РОЖДЕНИЯ}}": data.get("seller_birth_date", ""),
            "{{ПРОДАВЕЦ_АДРЕС}}":         n(data.get("seller_address", "")),
            "{{ПРОДАВЕЦ_ИНИЦИАЛЫ}}":      n(data.get("seller_initials", "")),
            "{{ПРОДАВЕЦ_ПОЛНЫЕ_ДАННЫЕ}}": n(data.get("seller_full_details", data.get("seller_name", ""))),

            # Идентификационная карта продавца (КР)
            "{{ПРОДАВЕЦ_ID}}":        data.get("seller_id_number", data.get("seller_id", "")),
            "{{ПРОДАВЕЦ_ID_НОМЕР}}":  data.get("seller_id_number", data.get("seller_id", "")),
            "{{ПРОДАВЕЦ_ID_ВЫДАНА}}": n(data.get("seller_id_issued_by", "")),
            "{{ПРОДАВЕЦ_ID_ДАТА}}":   data.get("seller_id_issued_date", ""),

            # Авто
            "{{МАРКА_МОДЕЛЬ}}": n(data.get("car_model", "")),
            "{{VIN}}":          data.get("car_vin", ""),
            "{{ГОД_ВЫП}}":      data.get("car_year", ""),
            "{{ЦВЕТ}}":         data.get("car_color", ""),
            "{{НОМ_КУЗОВА}}":   data.get("car_body_number", data.get("car_vin", "")),
            "{{НОМ_ТПО}}":      data.get("tpo_number", ""),
            "{{ДЕНЬ_ТПО}}":     data.get("tpo_day", ""),
            "{{МЕС_ТПО}}":      data.get("tpo_month", ""),
            "{{ГОД_ТПО}}":      data.get("tpo_year", ""),

            # Цена и оплата
            "{{ЦЕНА_ЦИФРАМИ}}":            price_fmt,
            "{{ЦЕНА_ПРОПИСЬЮ}}":           data.get("car_price_words", ""),
            "{{ВАЛЮТА}}":                  data.get("currency", "рублей"),
            # Фактические — расписка, акт, отчёт (документы дня расчёта)
            "{{СУММА_НАЛИЧНЫМИ}}":          cash_fmt,
            "{{СУММА_НАЛИЧНЫМИ_ПРОПИСЬЮ}}": cash_words,
            "{{КУРС_ФАКТИЧЕСКИЙ}}":         data.get("Фактический курс")
                                            or data.get("exchange_rate_fact", ""),
            # Расчётные — только поручение (Приложение № 1 к агентскому договору)
            "{{СУММА_ПОРУЧЕНИЯ}}":          order_fmt,
            "{{СУММА_ПОРУЧЕНИЯ_ПРОПИСЬЮ}}": order_words,
            "{{КУРС_ПОРУЧЕНИЯ}}":           data.get("exchange_rate", ""),
            # Старые имена — на случай не обновлённого шаблона
            "{{СУММА_ПРОПИСЬЮ}}":           order_words or cash_words,
            "{{КУРС_ДОЛЛАРА}}":             data.get("exchange_rate", ""),

            "{{ВАЛЮТА_НАЛИЧНЫМИ}}":        data.get("cash_currency", data.get("currency", "рублей")),

            # Реквизиты счёта — плейсхолдеры новой схемы ({{БАНК}}, {{СЧЕТ}} …)
            # добавляются ниже целиком из bank_requisites. Здесь только имена,
            # которые остались в шаблонах договоров: их текст пока не трогаем,
            # чтобы уже выпускаемые документы не изменились.
            "{{БАНК_КОРР_СТРОКА1}}": bank.get("corr_bank_name", ""),
            "{{БАНК_КОРР_СТРОКА2}}": bank.get("corr_bank_bic", ""),
            "{{БАНК_КОРР_СТРОКА3}}": bank.get("corr_bank_acc", ""),
            "{{БАНК_ПОЛ_СТРОКА1}}":  bank.get("bank_name", ""),
            # В договоре БИК и корр. счёт банка получателя стоят одной строкой.
            # В счёте мы их разнесли, договор оставлен как был.
            "{{БАНК_ПОЛ_СТРОКА2}}":  (f"БИК: {bank['bank_bic']}, корр. счёт: "
                                      f"{bank['bank_corr_acc']}")
                                     if bank.get("bank_bic") else "",
            "{{БАНК_ПРЯМОЙ_НАЗВАНИЕ}}": bank.get("bank_name", ""),
            "{{БАНК_ПРЯМОЙ_БИК}}":      bank.get("bank_bic", ""),
            "{{БАНК_ПРЯМОЙ_КОРР}}":     bank.get("bank_corr_acc", ""),
            "{{СЧЕТ_НОМЕР}}":           bank.get("account_number", ""),
        }

        # Карточка компании и реквизиты счёта — целиком, под своими именами.
        replacements.update(comp)
        replacements.update(br.placeholders(data))

        # ── Номер и дата в разрезе каждого документа ──────────────────────
        # Номер сделки один на весь комплект (присваивается один раз при
        # открытии сделки) — поэтому во все *_НОМЕР идёт одно и то же значение.
        # Даты по умолчанию равны дате сделки; конкретный build_* перекрывает
        # нужные через extra_replacements (например дату акта или отчёта).
        for scope in ("ДКП", "ДОГОВОРА", "СЧЕТА", "АКТА", "ОТЧЕТА"):
            replacements[f"{{{{НОМЕР_{scope}}}}}"] = number
            replacements[f"{{{{ДЕНЬ_{scope}}}}}"]  = day
            replacements[f"{{{{МЕСЯЦ_{scope}}}}}"] = self._month_name(month)
            replacements[f"{{{{ГОД_{scope}}}}}"]   = year

        # Дата ДКП — единственная дата комплекта, которая может отличаться от
        # даты сделки: ДКП нередко подписывают раньше агентского договора.
        # Живёт в колонке журнала «Дата ДКП» (в data при создании сделки —
        # ключ dkp_date). Пусто → остаётся дата договора, выставленная выше.
        # Подставляется здесь, а не в build_dkp, потому что на неё ссылаются
        # ещё акт и отчёт агента — иначе комплект разойдётся.
        dkp_date = str(data.get("Дата ДКП") or data.get("dkp_date") or "").strip()
        if len(dkp_date) >= 10:
            k_day, k_month, k_year = self._date_parts(dkp_date)
            replacements["{{ДЕНЬ_ДКП}}"]  = k_day
            replacements["{{МЕСЯЦ_ДКП}}"] = k_month
            replacements["{{ГОД_ДКП}}"]   = k_year

        # Номер ДКП — не номер сделки, а последние 6 знаков VIN.
        # Причина: номер сделки кодирует дату её открытия, а ДКП нередко
        # подписан раньше агентского договора — номер вида «110826001» на
        # договоре от 27 июля расшифровывается в неверную дату, и это видно
        # по всему комплекту. Номер по VIN даты не содержит.
        # Колонка журнала «Номер ДКП» перекрывает расчёт (ручной ввод).
        replacements["{{НОМЕР_ДКП}}"] = dkp_number_from(data, number)

        # Даты одним полем (ДД.ММ.ГГГГ) — используются в отчёте агента.
        # Приоритет: явно переданное значение (build_report) → колонка
        # журнала (ручной ввод) → пусто.
        replacements["{{ДАТА_ПОСТУПЛЕНИЯ}}"] = (
            data.get("payment_received_date")
            or data.get("Дата поступления")
            or ""
        )
        replacements["{{ДАТА_РАСЧЕТА}}"] = (
            data.get("settlement_date")
            or data.get("Дата расчёта")
            or ""
        )

        # Пробрасываемые снаружи плейсхолдеры (для актов и подобных документов
        # с расширенным набором). Перекрывают базовые при совпадении ключей.
        if extra_replacements:
            replacements.update(extra_replacements)

        # ── Какие плейсхолдеры реально есть в шаблоне ─────────────────────
        # Нужно ДО подстановки: пустое значение после замены неотличимо от
        # текста, который в шаблоне и должен быть пустым.
        present = self._scan_placeholders(doc)

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def replace_in_para(para):
            p_elem = para._element
            runs   = p_elem.findall(f"{{{W}}}r")
            if not runs:
                return

            full_text = "".join(
                t.text or ""
                for r in runs
                for t in r.findall(f"{{{W}}}t")
            )
            if not any(ph in full_text for ph in replacements):
                return

            # Шаг 1: заменяем плейсхолдеры, которые ПОЛНОСТЬЮ находятся
            # внутри одного run — сохраняем форматирование этого run и
            # НЕ трогаем соседние runs (например жирный номер пункта "2.", "4.", "6.").
            for r in runs:
                r_text = "".join(t.text or "" for t in r.findall(f"{{{W}}}t"))
                if any(ph in r_text for ph in replacements):
                    for t in r.findall(f"{{{W}}}t"):
                        if t.text:
                            new_text = t.text
                            for ph, val in replacements.items():
                                new_text = new_text.replace(ph, str(val) if val is not None else "")
                            t.text = new_text
                            if new_text and (new_text[0] == " " or new_text[-1] == " "):
                                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

            # Шаг 2: проверяем, остались ли НЕзамещённые плейсхолдеры —
            # значит они были разорваны между несколькими runs (например
            # «{{ДЕНЬ}}» — кавычка-ёлочка в одном run, а сам плейсхолдер
            # в другом). Для них делаем фоллбэк-слияние всех runs параграфа.
            runs = p_elem.findall(f"{{{W}}}r")
            full_text2 = "".join(
                t.text or ""
                for r in runs
                for t in r.findall(f"{{{W}}}t")
            )
            if not any(ph in full_text2 for ph in replacements):
                return

            # Фоллбэк: сливаем все runs параграфа в один (форматирование первого run).
            first_rpr  = runs[0].find(f"{{{W}}}rPr")
            children   = list(p_elem)
            insert_idx = children.index(runs[0])

            new_text = full_text2
            for ph, val in replacements.items():
                new_text = new_text.replace(ph, str(val) if val is not None else "")

            for r in runs:
                p_elem.remove(r)

            new_run = etree.Element(f"{{{W}}}r")
            if first_rpr is not None:
                new_run.append(deepcopy(first_rpr))
            new_t = etree.SubElement(new_run, f"{{{W}}}t")
            new_t.text = new_text
            if new_text and (new_text[0] == " " or new_text[-1] == " "):
                new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

            p_elem.insert(insert_idx, new_run)

        for para in doc.paragraphs:
            replace_in_para(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_para(para)

        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_in_para(para)
            for para in section.footer.paragraphs:
                replace_in_para(para)

        # ── Контроль заполнения ────────────────────────────────────────
        # 1) Незамещённые {{...}} — плейсхолдер есть в шаблоне, но его нет
        #    в карте замен (опечатка в шаблоне или забыли добавить в код).
        # 2) Пустые значения — плейсхолдер есть в шаблоне и он известен,
        #    но данных по нему в журнале нет. Именно так в документ уходит
        #    сторона без номера паспорта.
        leftover = sorted(self._scan_placeholders(doc))

        missing = sorted(
            PLACEHOLDER_LABELS.get(ph, ph.strip("{}"))
            for ph in present
            if ph in replacements
            and ph not in ALLOW_EMPTY_PLACEHOLDERS
            and not str(replacements[ph] or "").strip()
        )

        if leftover or missing:
            msg = (
                f"{output_name}.docx: незамещённые плейсхолдеры={leftover}, "
                f"пустые поля={missing}"
            )
            if STRICT_PLACEHOLDERS:
                logger.error(msg)
                raise MissingDataError(output_name, missing, leftover)
            logger.warning(msg + " (STRICT_PLACEHOLDERS=0 — документ выдан как есть)")

        # ── Удаляем w:proofErr (артефакты проверки правописания Word) ──
        # Эти теги между runs иногда вызывают переупорядочивание текста
        # при экспорте в PDF через LibreOffice. Header/footer — отдельные
        # XML-части, проходим по ним отдельно.
        for elem in doc.element.body.iter(f"{{{W}}}proofErr"):
            elem.getparent().remove(elem)
        for section in doc.sections:
            for elem in list(section.header._element.iter(f"{{{W}}}proofErr")):
                elem.getparent().remove(elem)
            for elem in list(section.footer._element.iter(f"{{{W}}}proofErr")):
                elem.getparent().remove(elem)

        path = self.output_dir / f"{output_name}.docx"
        doc.save(str(path))
        return str(path)
